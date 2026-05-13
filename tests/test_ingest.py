"""Tests for scripts/ingest.py.

Covers:
- scan_docx_dir: filename pattern filtering
- parse_all: chunk collection, --only filter, parse error resilience
- embed_all: delegates to embedder.embed_documents with correct args
- upsert_all: delegates to store.upsert_chunks
- CLI via typer.testing.CliRunner:
    - --dry-run: no embed/upsert calls
    - --reset: store.reset called before parsing
    - --only: only matching specs are parsed
    - missing docx dir: non-zero exit
"""

from __future__ import annotations

import sys
from pathlib import Path
from unittest.mock import MagicMock, call, patch

import pytest
from typer.testing import CliRunner

# ---------------------------------------------------------------------------
# Make src/ importable from the test runner's CWD
# ---------------------------------------------------------------------------
_REPO_ROOT = Path(__file__).parent.parent
if str(_REPO_ROOT / "src") not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT / "src"))

from spec_qa.parser import Chunk  # noqa: E402

# Import the script module
sys.path.insert(0, str(_REPO_ROOT / "scripts"))
from ingest import app, parse_all, embed_all, scan_docx_dir, upsert_all  # noqa: E402

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

runner = CliRunner()


def _make_chunk(
    chunk_id: str = "38.331-v18.5.0-5.1-0",
    spec_no: str = "38.331",
    version: str = "18.5.0",
    token_count: int = 10,
) -> Chunk:
    return Chunk(
        spec_no=spec_no,
        version=version,
        section_no="5.1",
        section_title="Introduction",
        heading_path=["5", "5.1"],
        text="Sample body text.",
        token_count=token_count,
        chunk_id=chunk_id,
    )


def _make_docx(tmp_path: Path, name: str) -> Path:
    """Create a minimal real docx file at tmp_path / name."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_heading("1 Introduction", level=1)
    doc.add_paragraph("Body text for testing.")
    p = tmp_path / name
    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# scan_docx_dir
# ---------------------------------------------------------------------------


class TestScanDocxDir:
    def test_finds_standard_filename(self, tmp_path: Path) -> None:
        _make_docx(tmp_path, "38.331-v18.5.0.docx")
        result = scan_docx_dir(tmp_path)
        assert len(result) == 1
        assert result[0].name == "38.331-v18.5.0.docx"

    def test_finds_hyphenated_spec(self, tmp_path: Path) -> None:
        _make_docx(tmp_path, "38.101-1-v19.5.0.docx")
        result = scan_docx_dir(tmp_path)
        assert len(result) == 1

    def test_ignores_non_matching_files(self, tmp_path: Path) -> None:
        # create a file that doesn't match the pattern
        (tmp_path / "readme.txt").touch()
        (tmp_path / "random.docx").touch()
        result = scan_docx_dir(tmp_path)
        assert result == []

    def test_returns_sorted(self, tmp_path: Path) -> None:
        _make_docx(tmp_path, "38.331-v18.5.0.docx")
        _make_docx(tmp_path, "38.211-v18.5.0.docx")
        result = scan_docx_dir(tmp_path)
        names = [p.name for p in result]
        assert names == sorted(names)

    def test_empty_dir(self, tmp_path: Path) -> None:
        assert scan_docx_dir(tmp_path) == []


# ---------------------------------------------------------------------------
# parse_all
# ---------------------------------------------------------------------------


class TestParseAll:
    def test_collects_chunks(self, tmp_path: Path) -> None:
        p = _make_docx(tmp_path, "38.331-v18.5.0.docx")
        chunks, failed = parse_all([p])
        assert len(chunks) > 0
        assert failed == []

    def test_only_filter_includes_matching(self, tmp_path: Path) -> None:
        p1 = _make_docx(tmp_path, "38.331-v18.5.0.docx")
        p2 = _make_docx(tmp_path, "38.211-v18.5.0.docx")
        chunks, failed = parse_all([p1, p2], only=["38.331"])
        spec_nos = {c.spec_no for c in chunks}
        assert "38.331" in spec_nos
        assert "38.211" not in spec_nos

    def test_only_filter_excludes_all(self, tmp_path: Path) -> None:
        p = _make_docx(tmp_path, "38.331-v18.5.0.docx")
        chunks, failed = parse_all([p], only=["99.999"])
        assert chunks == []
        assert failed == []

    def test_parse_error_skips_file(self, tmp_path: Path) -> None:
        # Pass a path that doesn't exist — should be caught and skipped
        bad_path = tmp_path / "38.999-v1.0.0.docx"
        good_path = _make_docx(tmp_path, "38.331-v18.5.0.docx")
        chunks, failed = parse_all([bad_path, good_path])
        # good file produces chunks; bad file is in failed
        assert len(chunks) > 0
        assert bad_path.name in failed

    def test_multiple_files_combined(self, tmp_path: Path) -> None:
        p1 = _make_docx(tmp_path, "38.331-v18.5.0.docx")
        p2 = _make_docx(tmp_path, "38.211-v18.5.0.docx")
        chunks, failed = parse_all([p1, p2])
        spec_nos = {c.spec_no for c in chunks}
        assert "38.331" in spec_nos
        assert "38.211" in spec_nos
        assert failed == []


# ---------------------------------------------------------------------------
# embed_all
# ---------------------------------------------------------------------------


class TestEmbedAll:
    def test_calls_embed_documents_with_texts(self) -> None:
        chunks = [
            _make_chunk("id-1", spec_no="38.331"),
            _make_chunk("id-2", spec_no="38.211"),
        ]
        mock_embedder = MagicMock()
        mock_embedder.embed_documents.return_value = [[0.1] * 4, [0.2] * 4]

        result = embed_all(chunks, mock_embedder, batch_size=64)

        mock_embedder.embed_documents.assert_called_once_with(
            ["Sample body text.", "Sample body text."],
            batch_size=64,
        )
        assert len(result) == 2

    def test_passes_batch_size(self) -> None:
        chunks = [_make_chunk()]
        mock_embedder = MagicMock()
        mock_embedder.embed_documents.return_value = [[0.0] * 4]

        embed_all(chunks, mock_embedder, batch_size=32)

        _, kwargs = mock_embedder.embed_documents.call_args
        assert kwargs["batch_size"] == 32

    def test_returns_embeddings_in_order(self) -> None:
        chunks = [_make_chunk("id-1"), _make_chunk("id-2"), _make_chunk("id-3")]
        expected = [[1.0], [2.0], [3.0]]
        mock_embedder = MagicMock()
        mock_embedder.embed_documents.return_value = expected

        result = embed_all(chunks, mock_embedder)
        assert result == expected


# ---------------------------------------------------------------------------
# upsert_all
# ---------------------------------------------------------------------------


class TestUpsertAll:
    def test_delegates_to_store(self) -> None:
        chunks = [_make_chunk("id-1"), _make_chunk("id-2")]
        embeddings = [[0.1] * 4, [0.2] * 4]
        mock_store = MagicMock()

        upsert_all(chunks, embeddings, mock_store)

        mock_store.upsert_chunks.assert_called_once_with(chunks, embeddings)


# ---------------------------------------------------------------------------
# CLI — dry-run
# ---------------------------------------------------------------------------


class TestCliDryRun:
    def test_dry_run_skips_embed_and_upsert(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "docx"
        docx_dir.mkdir()
        _make_docx(docx_dir, "38.331-v18.5.0.docx")

        with (
            patch("ingest.VoyageEmbedder") as mock_embedder_cls,
            patch("ingest.ChromaSpecStore") as mock_store_cls,
        ):
            mock_store = MagicMock()
            mock_store.count.return_value = 0
            mock_store_cls.return_value = mock_store

            result = runner.invoke(
                app,
                ["--docx-dir", str(docx_dir), "--dry-run"],
            )

        assert result.exit_code == 0, result.output
        # VoyageEmbedder should NOT be instantiated in dry-run mode
        mock_embedder_cls.assert_not_called()
        # upsert_chunks should NOT be called
        mock_store.upsert_chunks.assert_not_called()

    def test_dry_run_shows_chunk_count(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "docx"
        docx_dir.mkdir()
        _make_docx(docx_dir, "38.331-v18.5.0.docx")

        with (
            patch("ingest.VoyageEmbedder"),
            patch("ingest.ChromaSpecStore") as mock_store_cls,
        ):
            mock_store = MagicMock()
            mock_store.count.return_value = 0
            mock_store_cls.return_value = mock_store

            result = runner.invoke(
                app,
                ["--docx-dir", str(docx_dir), "--dry-run"],
            )

        assert "dry-run" in result.output.lower() or "Total" in result.output


# ---------------------------------------------------------------------------
# CLI — --reset
# ---------------------------------------------------------------------------


class TestCliReset:
    def test_reset_calls_store_reset(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "docx"
        docx_dir.mkdir()
        _make_docx(docx_dir, "38.331-v18.5.0.docx")

        with (
            patch("ingest.VoyageEmbedder") as mock_embedder_cls,
            patch("ingest.ChromaSpecStore") as mock_store_cls,
        ):
            mock_embedder = MagicMock()
            mock_embedder.embed_documents.return_value = [[0.1] * 4]
            mock_embedder_cls.return_value = mock_embedder

            mock_store = MagicMock()
            mock_store.count.return_value = 5
            mock_store_cls.return_value = mock_store

            result = runner.invoke(
                app,
                ["--docx-dir", str(docx_dir), "--reset", "--yes"],
            )

        assert result.exit_code == 0, result.output
        mock_store.reset.assert_called_once()

    def test_reset_without_yes_prompts(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "docx"
        docx_dir.mkdir()
        _make_docx(docx_dir, "38.331-v18.5.0.docx")

        with (
            patch("ingest.VoyageEmbedder"),
            patch("ingest.ChromaSpecStore") as mock_store_cls,
        ):
            mock_store = MagicMock()
            mock_store.count.return_value = 3
            mock_store_cls.return_value = mock_store

            # Provide "n" to cancel the prompt
            result = runner.invoke(
                app,
                ["--docx-dir", str(docx_dir), "--reset"],
                input="n\n",
            )

        # Aborted — reset should NOT have been called
        mock_store.reset.assert_not_called()

    def test_no_reset_does_not_call_store_reset(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "docx"
        docx_dir.mkdir()
        _make_docx(docx_dir, "38.331-v18.5.0.docx")

        with (
            patch("ingest.VoyageEmbedder") as mock_embedder_cls,
            patch("ingest.ChromaSpecStore") as mock_store_cls,
        ):
            mock_embedder = MagicMock()
            mock_embedder.embed_documents.return_value = [[0.1] * 4]
            mock_embedder_cls.return_value = mock_embedder

            mock_store = MagicMock()
            mock_store.count.return_value = 0
            mock_store_cls.return_value = mock_store

            result = runner.invoke(
                app,
                ["--docx-dir", str(docx_dir)],
            )

        mock_store.reset.assert_not_called()


# ---------------------------------------------------------------------------
# CLI — --only filter
# ---------------------------------------------------------------------------


class TestCliOnly:
    def test_only_filters_to_spec(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "docx"
        docx_dir.mkdir()
        _make_docx(docx_dir, "38.331-v18.5.0.docx")
        _make_docx(docx_dir, "38.211-v18.5.0.docx")

        captured_chunks: list[list[Chunk]] = []

        original_upsert = upsert_all

        def _spy_upsert(chunks, embeddings, store):
            captured_chunks.append(chunks)
            original_upsert(chunks, embeddings, store)

        with (
            patch("ingest.VoyageEmbedder") as mock_embedder_cls,
            patch("ingest.ChromaSpecStore") as mock_store_cls,
            patch("ingest.upsert_all", side_effect=_spy_upsert),
        ):
            mock_embedder = MagicMock()
            # embed_documents returns one vec per chunk text
            mock_embedder.embed_documents.side_effect = lambda texts, **kw: [
                [0.1] * 4 for _ in texts
            ]
            mock_embedder_cls.return_value = mock_embedder

            mock_store = MagicMock()
            mock_store.count.return_value = 1
            mock_store_cls.return_value = mock_store

            result = runner.invoke(
                app,
                ["--docx-dir", str(docx_dir), "--only", "38.331"],
            )

        assert result.exit_code == 0, result.output
        if captured_chunks:
            all_spec_nos = {c.spec_no for chunks in captured_chunks for c in chunks}
            assert "38.331" in all_spec_nos
            assert "38.211" not in all_spec_nos


# ---------------------------------------------------------------------------
# CLI — missing docx dir
# ---------------------------------------------------------------------------


class TestCliMissingDir:
    def test_exits_nonzero_when_dir_empty(self, tmp_path: Path) -> None:
        docx_dir = tmp_path / "empty_dir"
        docx_dir.mkdir()

        with patch("ingest.ChromaSpecStore"):
            result = runner.invoke(app, ["--docx-dir", str(docx_dir)])

        assert result.exit_code != 0
