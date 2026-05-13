"""Tests for src/spec_qa/parser.py.

Uses unittest.mock to create lightweight docx-like objects so that no real
.docx files are required for the unit tests.
"""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest

from spec_qa.parser import (
    MAX_TOKENS_PER_CHUNK,
    Chunk,
    Paragraph,
    _SectionState,
    _extract_title,
    _parse_filename,
    _parse_heading_text,
    chunk_by_section,
    count_tokens,
    extract_paragraphs,
    parse_spec_file,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_para(text: str, style_name: str = "Normal") -> MagicMock:
    """Build a minimal mock of a python-docx Paragraph."""
    para = MagicMock()
    para.text = text
    para.style.name = style_name
    return para


def _make_docx_table(rows_texts: list[list[str]]) -> MagicMock:
    """Build a mock Table where each row is a list of cell texts."""
    from docx.table import Table  # only to satisfy isinstance checks in parser

    table = MagicMock(spec=Table)
    mock_rows = []
    for row_texts in rows_texts:
        row = MagicMock()
        mock_cells = []
        for cell_text in row_texts:
            cell = MagicMock()
            cell_para = _make_docx_para(cell_text)
            cell.paragraphs = [cell_para]
            mock_cells.append(cell)
        row.cells = mock_cells
        mock_rows.append(row)
    table.rows = mock_rows
    return table


# ---------------------------------------------------------------------------
# _parse_heading_text
# ---------------------------------------------------------------------------


class TestParseHeadingText:
    def test_single_level(self):
        result = _parse_heading_text("5 General")
        assert result == (1, "5")

    def test_two_levels(self):
        result = _parse_heading_text("5.2 Procedures")
        assert result == (2, "5.2")

    def test_four_levels(self):
        result = _parse_heading_text("5.3.5.4 RRC Connection Reconfiguration")
        assert result == (4, "5.3.5.4")

    def test_seven_levels(self):
        result = _parse_heading_text("5.3.5.4.1.2.3 Deep heading")
        assert result == (7, "5.3.5.4.1.2.3")

    def test_annex_top_level(self):
        result = _parse_heading_text("Annex A Change history")
        assert result is not None
        level, section_no = result
        assert level == 1
        assert section_no == "Annex A"

    def test_annex_sub_level(self):
        result = _parse_heading_text("Annex B.1 Some Annex")
        assert result is not None
        level, section_no = result
        assert level == 2
        assert section_no == "Annex B.1"

    def test_annex_two_sub_levels(self):
        result = _parse_heading_text("Annex C.2.3 Deep annex")
        assert result is not None
        level, section_no = result
        assert level == 3
        assert section_no == "Annex C.2.3"

    def test_body_text_not_heading(self):
        assert _parse_heading_text("This is just a body paragraph.") is None

    def test_number_only_not_heading(self):
        # A bare number with no following text should not match
        # (regex requires \s+\S after the number)
        assert _parse_heading_text("5") is None

    def test_table_reference_not_heading(self):
        # "Table 5.3.1:" starts with a word, not a digit
        assert _parse_heading_text("Table 5.3.1: Parameter mapping") is None


# ---------------------------------------------------------------------------
# _extract_title
# ---------------------------------------------------------------------------


class TestExtractTitle:
    def test_numeric(self):
        assert _extract_title("5.3.5.4 RRC Connection Reconfiguration", "5.3.5.4") == \
               "RRC Connection Reconfiguration"

    def test_annex(self):
        assert _extract_title("Annex B.1 Some Annex", "Annex B.1") == "Some Annex"

    def test_no_section_no(self):
        assert _extract_title("Foreword", "") == "Foreword"


# ---------------------------------------------------------------------------
# _parse_filename
# ---------------------------------------------------------------------------


class TestParseFilename:
    def test_simple(self):
        assert _parse_filename(Path("38.331-v18.5.0.docx")) == ("38.331", "18.5.0")

    def test_hyphenated_spec(self):
        assert _parse_filename(Path("38.101-1-v19.5.0.docx")) == ("38.101-1", "19.5.0")

    def test_double_hyphen_spec(self):
        assert _parse_filename(Path("36.331-2-v16.0.0.docx")) == ("36.331-2", "16.0.0")

    def test_invalid_raises(self):
        with pytest.raises(ValueError):
            _parse_filename(Path("no_version_here.docx"))


# ---------------------------------------------------------------------------
# count_tokens
# ---------------------------------------------------------------------------


class TestCountTokens:
    def test_empty(self):
        assert count_tokens("") == 0

    def test_positive(self):
        assert count_tokens("Hello world") > 0

    def test_longer_is_more(self):
        short = count_tokens("Hello")
        long = count_tokens("Hello world, this is a longer sentence with more tokens.")
        assert long > short


# ---------------------------------------------------------------------------
# extract_paragraphs (mocked docx)
# ---------------------------------------------------------------------------


def _fake_block_items(doc):
    """Yield mock paragraphs for a tiny synthetic document."""
    yield _make_docx_para("1 Introduction", "Heading 1")
    yield _make_docx_para("This is the introduction body.")
    yield _make_docx_para("")  # should be skipped
    yield _make_docx_para("1.1 Background", "Heading 2")
    yield _make_docx_para("Background content here.")


class TestExtractParagraphs:
    def test_empty_docx(self, tmp_path):
        """A docx with no paragraphs should return an empty list."""
        from docx import Document as DocxDocument

        p = tmp_path / "empty.docx"
        DocxDocument().save(str(p))
        result = extract_paragraphs(p)
        assert result == []

    def test_skips_empty_paragraphs(self, tmp_path):
        """Empty paragraphs must be filtered out."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("")
        doc.add_paragraph("Hello world")
        doc.add_paragraph("")
        p = tmp_path / "doc.docx"
        doc.save(str(p))
        result = extract_paragraphs(p)
        assert len(result) == 1
        assert result[0].text == "Hello world"

    def test_heading_style_recognised(self, tmp_path):
        """A 'Heading 1' style paragraph must be recognised as heading level 1."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("1 Introduction", level=1)
        doc.add_paragraph("Body text.")
        p = tmp_path / "doc.docx"
        doc.save(str(p))
        result = extract_paragraphs(p)
        headings = [r for r in result if r.is_heading]
        assert len(headings) >= 1
        assert headings[0].heading_level == 1

    def test_text_heuristic_heading(self, tmp_path):
        """A paragraph starting with a numeric section number is a heading."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        # Add with Normal style so only text heuristic can detect it
        doc.add_paragraph("5.3.5.4 RRC Connection Reconfiguration")
        p = tmp_path / "doc.docx"
        doc.save(str(p))
        result = extract_paragraphs(p)
        assert len(result) == 1
        r = result[0]
        assert r.is_heading is True
        assert r.heading_level == 4
        assert r.section_no == "5.3.5.4"

    def test_annex_heading(self, tmp_path):
        """A paragraph starting with 'Annex B.1' is an annex heading."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Annex B.1 Some Annex")
        p = tmp_path / "doc.docx"
        doc.save(str(p))
        result = extract_paragraphs(p)
        assert len(result) == 1
        r = result[0]
        assert r.is_heading is True
        assert r.section_no == "Annex B.1"
        assert r.heading_level == 2


# ---------------------------------------------------------------------------
# chunk_by_section
# ---------------------------------------------------------------------------


def _paras(*items: tuple[str, bool, int | None, str]) -> list[Paragraph]:
    """Convenience builder: (text, is_heading, heading_level, section_no)."""
    return [
        Paragraph(text=t, style="", is_heading=h, heading_level=lvl, section_no=sno)
        for t, h, lvl, sno in items
    ]


class TestChunkBySection:
    def test_single_section(self):
        paras = _paras(
            ("1 Introduction", True, 1, "1"),
            ("Body paragraph one.", False, None, ""),
            ("Body paragraph two.", False, None, ""),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        assert len(chunks) == 1
        c = chunks[0]
        assert c.section_no == "1"
        assert c.spec_no == "38.331"
        assert c.version == "18.5.0"
        assert "Body paragraph one." in c.text
        assert "Body paragraph two." in c.text
        assert c.token_count > 0

    def test_multiple_sections(self):
        paras = _paras(
            ("1 Introduction", True, 1, "1"),
            ("Intro body.", False, None, ""),
            ("2 Scope", True, 1, "2"),
            ("Scope body.", False, None, ""),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        section_nos = [c.section_no for c in chunks]
        assert "1" in section_nos
        assert "2" in section_nos

    def test_heading_path_hierarchy(self):
        paras = _paras(
            ("5 General", True, 1, "5"),
            ("5.3 Procedures", True, 2, "5.3"),
            ("5.3.5 Sub-procedure", True, 3, "5.3.5"),
            ("5.3.5.4 RRC Connection Reconfiguration", True, 4, "5.3.5.4"),
            ("Body content.", False, None, ""),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        # The deepest section chunk should carry full hierarchy
        deep = next(c for c in chunks if c.section_no == "5.3.5.4")
        assert deep.heading_path == ["5", "5.3", "5.3.5", "5.3.5.4"]

    def test_chunk_id_format(self):
        paras = _paras(
            ("5.3.5.4 RRC Connection Reconfiguration", True, 4, "5.3.5.4"),
            ("Body.", False, None, ""),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        assert chunks[0].chunk_id.startswith("38.331-v18.5.0-5.3.5.4-")

    def test_empty_paragraphs(self):
        chunks = chunk_by_section([], "38.331", "18.5.0")
        assert chunks == []

    def test_no_heading_preamble(self):
        """Paragraphs before any heading land in section_no='' preamble."""
        paras = _paras(
            ("Copyright notice text.", False, None, ""),
            ("Foreword content.", False, None, ""),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        assert len(chunks) == 1
        assert chunks[0].section_no == ""

    def test_token_split(self):
        """A section with many large paragraphs must be split into multiple chunks."""
        # Create paragraphs each ~150 tokens; 10 of them ≈ 1500 tokens > MAX
        long_text = "word " * 150  # ~150 tokens per paragraph
        paras = _paras(
            ("5 General", True, 1, "5"),
            *((long_text, False, None, "") for _ in range(10)),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        section_chunks = [c for c in chunks if c.section_no == "5"]
        assert len(section_chunks) > 1, "Expected at least 2 chunks for large section"
        for c in section_chunks:
            assert c.token_count <= MAX_TOKENS_PER_CHUNK + count_tokens(long_text), \
                "Chunk token count should not exceed limit by more than one paragraph"

    def test_overlap_in_split_chunks(self):
        """The second split chunk must start with text from the end of the first."""
        long_text_a = "alpha " * 150
        long_text_b = "beta " * 150
        long_text_c = "gamma " * 150
        long_text_d = "delta " * 150
        long_text_e = "epsilon " * 150
        long_text_f = "zeta " * 150
        long_text_g = "eta " * 150
        long_text_h = "theta " * 150
        long_text_i = "iota " * 150

        paras = _paras(
            ("5 General", True, 1, "5"),
            (long_text_a, False, None, ""),
            (long_text_b, False, None, ""),
            (long_text_c, False, None, ""),
            (long_text_d, False, None, ""),
            (long_text_e, False, None, ""),
            (long_text_f, False, None, ""),
            (long_text_g, False, None, ""),
            (long_text_h, False, None, ""),
            (long_text_i, False, None, ""),
        )
        chunks = chunk_by_section(paras, "38.331", "18.5.0")
        section_chunks = [c for c in chunks if c.section_no == "5"]
        assert len(section_chunks) >= 2
        # The text that ended chunk[0] must appear in chunk[1] as overlap
        first_chunk_end = section_chunks[0].text.split("\n\n")[-1]
        assert first_chunk_end in section_chunks[1].text


# ---------------------------------------------------------------------------
# Chunk.to_dict / Chunk.metadata
# ---------------------------------------------------------------------------


class TestChunkMethods:
    def _make_chunk(self) -> Chunk:
        return Chunk(
            spec_no="38.331",
            version="18.5.0",
            section_no="5.3.5.4",
            section_title="RRC Connection Reconfiguration",
            heading_path=["5", "5.3", "5.3.5", "5.3.5.4"],
            text="Some body text.",
            token_count=4,
            chunk_id="38.331-v18.5.0-5.3.5.4-0",
        )

    def test_to_dict_contains_text(self):
        d = self._make_chunk().to_dict()
        assert d["text"] == "Some body text."
        assert d["chunk_id"] == "38.331-v18.5.0-5.3.5.4-0"

    def test_metadata_no_text(self):
        m = self._make_chunk().metadata()
        assert "text" not in m
        assert m["spec_no"] == "38.331"
        assert m["heading_path_str"] == "5 > 5.3 > 5.3.5 > 5.3.5.4"

    def test_metadata_all_keys(self):
        m = self._make_chunk().metadata()
        for key in ("chunk_id", "spec_no", "version", "section_no",
                    "section_title", "heading_path_str", "token_count"):
            assert key in m, f"Missing key: {key}"


# ---------------------------------------------------------------------------
# parse_spec_file (integration via tmp real docx)
# ---------------------------------------------------------------------------


class TestParseSpecFile:
    def test_roundtrip(self, tmp_path):
        """Create a minimal real docx, parse it, and verify chunk structure."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("1 Introduction", level=1)
        doc.add_paragraph("This is the introduction body.")
        doc.add_heading("2 Scope", level=1)
        doc.add_paragraph("This document covers NR radio specifications.")

        p = tmp_path / "38.331-v18.5.0.docx"
        doc.save(str(p))

        chunks = parse_spec_file(p)
        assert len(chunks) >= 1

        spec_nos = {c.spec_no for c in chunks}
        assert spec_nos == {"38.331"}

        versions = {c.version for c in chunks}
        assert versions == {"18.5.0"}

    def test_hyphenated_spec_filename(self, tmp_path):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_heading("1 Introduction", level=1)
        doc.add_paragraph("Body.")
        p = tmp_path / "38.101-1-v19.5.0.docx"
        doc.save(str(p))

        chunks = parse_spec_file(p)
        assert all(c.spec_no == "38.101-1" for c in chunks)
        assert all(c.version == "19.5.0" for c in chunks)
