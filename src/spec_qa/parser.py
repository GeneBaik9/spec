"""3GPP TS docx parser: extract paragraphs and chunk by section.

Supports hierarchical section numbers (e.g. "5.3.5.4") and Annex headings
(e.g. "Annex B.1"). Chunks are sized to MAX_TOKENS_PER_CHUNK with a small
overlap between adjacent chunks in the same section.
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import tiktoken
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph as DocxParagraph

# ---------------------------------------------------------------------------
# Token counting
# ---------------------------------------------------------------------------

_ENC = tiktoken.get_encoding("cl100k_base")

MAX_TOKENS_PER_CHUNK = 1200
OVERLAP_PARAGRAPHS = 2  # number of trailing paragraphs to prepend to next chunk


def count_tokens(text: str) -> int:
    """Return the number of cl100k_base tokens in *text*."""
    return len(_ENC.encode(text))


# ---------------------------------------------------------------------------
# Heading recognition regexes
# ---------------------------------------------------------------------------

# Matches "5", "5.2", "5.2.3", "5.2.3.1", …  followed by at least one non-ws char
_RE_NUMERIC_HEADING = re.compile(r"^(\d+(?:\.\d+)*)\s+\S")

# Matches "Annex A", "Annex B.1", "Annex C.2.3", … followed by at least one non-ws char
_RE_ANNEX_HEADING = re.compile(r"^(Annex\s+[A-Z](?:\.\d+)*)\s+\S", re.IGNORECASE)

# Extract just the section number from an annex string: "Annex B.1" → "B.1"
_RE_ANNEX_LETTER = re.compile(r"^Annex\s+([A-Z](?:\.\d+)*)", re.IGNORECASE)


def _heading_level_from_style(style_name: str) -> int | None:
    """Return heading level (1-based) from a 'Heading N' style name, else None."""
    m = re.match(r"^Heading\s+(\d+)$", style_name, re.IGNORECASE)
    return int(m.group(1)) if m else None


def _parse_heading_text(text: str) -> tuple[int, str] | None:
    """Infer (level, section_no) from paragraph text using regex heuristics.

    Returns None when the text does not look like a heading.
    """
    # Priority 1: numeric section number
    m = _RE_NUMERIC_HEADING.match(text)
    if m:
        section_no = m.group(1)
        level = section_no.count(".") + 1
        return level, section_no

    # Priority 2: Annex heading
    m = _RE_ANNEX_HEADING.match(text)
    if m:
        annex_str = m.group(1)  # e.g. "Annex B.1"
        letter_part_m = _RE_ANNEX_LETTER.match(annex_str)
        letter_part = letter_part_m.group(1) if letter_part_m else annex_str
        # "A" → level 1, "B.1" → level 2, "C.2.3" → level 3
        level = letter_part.count(".") + 1
        section_no = f"Annex {letter_part}"
        return level, section_no

    return None


# ---------------------------------------------------------------------------
# Paragraph dataclass
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class Paragraph:
    """Normalized paragraph from a docx document."""

    text: str
    style: str
    is_heading: bool
    heading_level: int | None  # 1-based; None for body paragraphs
    section_no: str  # set when is_heading=True, else ""


# ---------------------------------------------------------------------------
# Chunk dataclass
# ---------------------------------------------------------------------------


@dataclass
class Chunk:
    """A semantically coherent text chunk from a 3GPP spec section."""

    spec_no: str          # e.g. "38.331"
    version: str          # e.g. "18.5.0"
    section_no: str       # e.g. "5.3.5.4" or "Annex B.1" or "" (preamble)
    section_title: str    # e.g. "RRC Connection Reconfiguration"
    heading_path: list[str] = field(default_factory=list)  # ["5", "5.3", "5.3.5", "5.3.5.4"]
    text: str = ""
    token_count: int = 0
    chunk_id: str = ""

    # ------------------------------------------------------------------
    def to_dict(self) -> dict:
        """Flat dict representation suitable for serialisation."""
        return {
            "chunk_id": self.chunk_id,
            "spec_no": self.spec_no,
            "version": self.version,
            "section_no": self.section_no,
            "section_title": self.section_title,
            "heading_path": self.heading_path,
            "text": self.text,
            "token_count": self.token_count,
        }

    def metadata(self) -> dict:
        """Return metadata dict (everything except *text*) for a vector store."""
        return {
            "chunk_id": self.chunk_id,
            "spec_no": self.spec_no,
            "version": self.version,
            "section_no": self.section_no,
            "section_title": self.section_title,
            "heading_path_str": " > ".join(self.heading_path),
            "token_count": self.token_count,
        }


# ---------------------------------------------------------------------------
# Paragraph extraction
# ---------------------------------------------------------------------------


def _iter_block_items(document: Document):
    """Yield paragraphs and tables in document body order.

    python-docx's document.paragraphs skips table cells, so we walk the raw
    XML to preserve reading order.
    """
    from docx.oxml.ns import qn  # local import to avoid polluting namespace

    body = document.element.body
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            yield DocxParagraph(child, document)
        elif tag == "tbl":
            yield Table(child, document)


def extract_paragraphs(docx_path: Path) -> list[Paragraph]:
    """Open *docx_path* and return normalised Paragraph list in document order.

    Both body paragraphs and table-cell paragraphs are included.
    Empty paragraphs are skipped.
    """
    doc = Document(str(docx_path))
    results: list[Paragraph] = []

    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            # Iterate all cells row-by-row; multi-column tables are serialised
            # in row-major order, which may not match visual reading order.
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_paragraph(para, results)
        else:
            _process_paragraph(block, results)

    return results


def _process_paragraph(para: DocxParagraph, out: list[Paragraph]) -> None:
    """Normalise one python-docx paragraph and append to *out* if non-empty."""
    text = para.text.strip()
    if not text:
        return

    style_name: str = para.style.name if para.style else ""

    # Priority 1: explicit Heading style
    level = _heading_level_from_style(style_name)
    if level is not None:
        # Extract section_no from text when it starts with a number / Annex
        parsed = _parse_heading_text(text)
        section_no = parsed[1] if parsed else ""
        out.append(Paragraph(
            text=text,
            style=style_name,
            is_heading=True,
            heading_level=level,
            section_no=section_no,
        ))
        return

    # Priority 2: text-based heuristic
    parsed = _parse_heading_text(text)
    if parsed is not None:
        level, section_no = parsed
        out.append(Paragraph(
            text=text,
            style=style_name,
            is_heading=True,
            heading_level=level,
            section_no=section_no,
        ))
        return

    # Body paragraph
    out.append(Paragraph(
        text=text,
        style=style_name,
        is_heading=False,
        heading_level=None,
        section_no="",
    ))


# ---------------------------------------------------------------------------
# Section state tracker
# ---------------------------------------------------------------------------


@dataclass
class _SectionState:
    """Tracks the current heading hierarchy while scanning paragraphs."""

    # Map from level → (section_no, title_text)
    # e.g. {1: ("5", "General"), 2: ("5.3", "Procedures"), ...}
    _stack: dict[int, tuple[str, str]] = field(default_factory=dict)

    def update(self, para: Paragraph) -> None:
        """Register a heading paragraph and prune deeper levels."""
        level = para.heading_level  # guaranteed non-None for headings
        assert level is not None

        # Derive title: everything after the section_no prefix
        title = _extract_title(para.text, para.section_no)

        self._stack[level] = (para.section_no, title)
        # Remove all deeper levels so they don't linger after a shallower heading
        for deeper in list(self._stack.keys()):
            if deeper > level:
                del self._stack[deeper]

    @property
    def current_section_no(self) -> str:
        if not self._stack:
            return ""
        return self._stack[max(self._stack)][0]

    @property
    def current_title(self) -> str:
        if not self._stack:
            return ""
        return self._stack[max(self._stack)][1]

    @property
    def heading_path(self) -> list[str]:
        """Ordered list of section numbers from root to current."""
        return [self._stack[lvl][0] for lvl in sorted(self._stack)]


def _extract_title(text: str, section_no: str) -> str:
    """Strip the section_no prefix from heading text to get the title."""
    if not section_no:
        return text.strip()
    # Remove the leading section number (and any trailing whitespace)
    # Handle both "5.3 Title" and "Annex B.1 Title"
    stripped = text[len(section_no):].strip()
    return stripped if stripped else text.strip()


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------


def _make_chunk_id(spec_no: str, version: str, section_no: str, idx: int) -> str:
    """Create a stable, filesystem-safe chunk identifier."""
    slug = re.sub(r"[^A-Za-z0-9._-]", "_", section_no) if section_no else "preamble"
    return f"{spec_no}-v{version}-{slug}-{idx}"


def chunk_by_section(
    paragraphs: list[Paragraph],
    spec_no: str,
    version: str,
) -> list[Chunk]:
    """Group *paragraphs* into Chunks, one per section, splitting on token limit.

    Sections longer than MAX_TOKENS_PER_CHUNK are split at paragraph boundaries
    with OVERLAP_PARAGRAPHS overlap between consecutive split chunks.
    """
    state = _SectionState()
    chunks: list[Chunk] = []

    # Buffer of body paragraphs for the current section
    body_buf: list[str] = []

    def _flush(section_no: str, section_title: str, heading_path: list[str]) -> None:
        """Emit chunks from *body_buf* for the completed section."""
        if not body_buf:
            return
        _split_and_emit(
            body_buf, spec_no, version, section_no, section_title, heading_path, chunks
        )
        body_buf.clear()

    for para in paragraphs:
        if para.is_heading:
            # Flush accumulated body before moving to new section
            _flush(state.current_section_no, state.current_title, state.heading_path)
            state.update(para)
        else:
            body_buf.append(para.text)

    # Flush the final section
    _flush(state.current_section_no, state.current_title, state.heading_path)

    return chunks


def _split_and_emit(
    body_texts: list[str],
    spec_no: str,
    version: str,
    section_no: str,
    section_title: str,
    heading_path: list[str],
    out: list[Chunk],
) -> None:
    """Emit one or more Chunks from *body_texts*, respecting MAX_TOKENS_PER_CHUNK.

    When splitting is needed, the last OVERLAP_PARAGRAPHS paragraphs of the
    previous chunk are prepended to the next chunk (overlap).
    """
    current_paras: list[str] = []
    current_tokens: int = 0
    idx: int = 0

    # Track how many chunks have been emitted for this section so far
    # (used to decide whether to include overlap)
    overlap_prefix: list[str] = []

    for para_text in body_texts:
        para_tokens = count_tokens(para_text)

        if current_tokens + para_tokens > MAX_TOKENS_PER_CHUNK and current_paras:
            # Emit current buffer as a chunk
            text = "\n\n".join(current_paras)
            out.append(Chunk(
                spec_no=spec_no,
                version=version,
                section_no=section_no,
                section_title=section_title,
                heading_path=list(heading_path),
                text=text,
                token_count=count_tokens(text),
                chunk_id=_make_chunk_id(spec_no, version, section_no, idx),
            ))

            # Prepare overlap: last OVERLAP_PARAGRAPHS from the emitted chunk
            overlap_prefix = current_paras[-OVERLAP_PARAGRAPHS:]
            idx += 1

            # Start new buffer with overlap prefix
            current_paras = list(overlap_prefix) + [para_text]
            current_tokens = sum(count_tokens(p) for p in current_paras)
        else:
            current_paras.append(para_text)
            current_tokens += para_tokens

    # Emit remainder
    if current_paras:
        text = "\n\n".join(current_paras)
        out.append(Chunk(
            spec_no=spec_no,
            version=version,
            section_no=section_no,
            section_title=section_title,
            heading_path=list(heading_path),
            text=text,
            token_count=count_tokens(text),
            chunk_id=_make_chunk_id(spec_no, version, section_no, idx),
        ))


# ---------------------------------------------------------------------------
# Filename parsing
# ---------------------------------------------------------------------------


def _parse_filename(docx_path: Path) -> tuple[str, str]:
    """Extract (spec_no, version) from a filename like '38.331-v18.5.0.docx'.

    Handles hyphenated spec numbers such as '38.101-1-v19.5.0.docx'
    → spec_no='38.101-1', version='19.5.0'.
    """
    stem = docx_path.stem  # e.g. "38.331-v18.5.0" or "38.101-1-v19.5.0"
    # The version tag is always "-vMAJOR.MINOR.PATCH" at the end
    m = re.search(r"-v(\d+\.\d+\.\d+)$", stem)
    if not m:
        raise ValueError(f"Cannot parse version from filename: {docx_path.name}")
    version = m.group(1)
    spec_no = stem[: m.start()]  # everything before "-vX.Y.Z"
    return spec_no, version


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def parse_spec_file(docx_path: Path) -> list[Chunk]:
    """Parse a 3GPP TS docx file and return all Chunks.

    The spec number and version are inferred from the filename convention
    ``{spec_no}-v{major}.{minor}.{patch}.docx``.
    """
    spec_no, version = _parse_filename(docx_path)
    paragraphs = extract_paragraphs(docx_path)
    return chunk_by_section(paragraphs, spec_no, version)


# ---------------------------------------------------------------------------
# CLI smoke-test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python parser.py <path-to-docx>")
        sys.exit(1)

    path = Path(sys.argv[1])
    chunks = parse_spec_file(path)
    print(f"Total chunks: {len(chunks)}")
    if chunks:
        first = chunks[0]
        print(f"\nFirst chunk:")
        print(f"  chunk_id    : {first.chunk_id}")
        print(f"  section_no  : {first.section_no!r}")
        print(f"  section_title: {first.section_title!r}")
        print(f"  heading_path: {first.heading_path}")
        print(f"  token_count : {first.token_count}")
        print(f"  text[:200]  : {first.text[:200]!r}")
